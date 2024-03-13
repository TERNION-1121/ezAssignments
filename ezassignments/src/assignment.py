from docx import Document
from docx.shared import Inches

from os import listdir
from os.path import join

from .question_classes import *
from .utils import *


class Assignment:
    def __init__(self, dir_path: str, txt_name: str, docx_name: str) -> None:
        """
        Initialize an Assignment object if `txt_name.txt` is found in the directory given at `dir_path`
        """
        self.dir_path = dir_path
        file_path = join(self.dir_path, txt_name)
        try:
            file = open(file_path, 'rt')
            self.docx_path = join(self.dir_path, docx_name)

        except FileNotFoundError:
            print(f"File {file_path} was not found.")
            return

        # strip whitespaces from each line if not a newline itself
        contents = [line.strip('\n').strip() if line != '\n' else '\n' for line in file]
        file.close()

        # strip newlines from the start and end of contents
        start = 0
        while contents[start] == '\n':
            start += 1
        end = len(contents) - 1
        while contents[end] == '\n':
            end -= 1

        # split the questions
        contents = split_contents(contents[start:end+1])
        # store the question classes
        self.QUESTION_CLS = (MCQ, AR, SUB)

        # process each question
        for question in contents:
            tag = question[0]  # type of question as in input
            for q_cls in self.QUESTION_CLS:  # each question class
                if q_cls.__name__ != tag:
                    continue
                # tag matches a question class
                if not q_cls.from_list_content(question[1:]):  # question did not fit the given question type
                    Invalid(question=question)  # add question to the Invalid class
                break
            else:  # question had no valid question type as tag
                Invalid(question=question)

        # check for images dir in dir_path
        path = join(self.dir_path, "images")
        try:
            self.image_assets = {f for f in listdir(path)}
        except FileNotFoundError:
            print(f"Warning: {path} was not found, image insertion won't be possible.", end="\n\n")
            self.image_assets = None

        self.doc = Document()
        self.doc_title = "Assignment"

    def process_title(self):
        yes_opts = ('y', 'yes')
        no_opts = ('n', 'no')
        # ask choice
        print(f"Change document title? [default is {self.doc_title}] (y/n):", end=" ")
        while (choice := input().strip().lower()) not in yes_opts + no_opts:
            print("\033[1A\033[0J", end='')
            # ^ ANSI escape codes reference: https://gist.github.com/fnky/458719343aabd01cfb17a3a4f7296797
            print(f"Change document title? [default is {self.doc_title}] (y/n):", end=" ")
        # user opted for default title
        if choice in no_opts:
            print()
            return
        # ask user for title
        while not (title := input("Write your own document title: ")):
            print("\033[1A\033[0J", end='')

        print()
        self.doc_title = title

    def attach_image(self, q_obj: Question):
        if not q_obj:
            return 
        if not self.image_assets:
            Invalid(image=q_obj.image)
            return 
        
        for image in self.image_assets:
            if image == q_obj.image:
                self.doc.add_picture(join(self.dir_path, "images", q_obj.image), width=Inches(6))
                self.doc.save(self.docx_path)
                break 
        else:  # image was not found in image_assets; invalidated  
            Invalid(image=q_obj.image)

    def init_doc(self):
        """
        Basic initializations for the .docx file
        """
        self.doc.add_heading(self.doc_title, level=0)
        self.doc.save(self.docx_path)

    def add_MCQS(self):
        """
        Add the valid Multiple Choice Questions to the .docx file
        """
        self.doc.add_heading("Multiple Choice Questions", level=1)
        # add instruction
        MCQ.write_instructions(self.doc)
        # add MCQ questions
        for mcq in MCQ.QUESTIONS:
            # attach image, if any
            self.attach_image(mcq)
            # add statement
            statement = ' '.join(mcq.question)
            self.doc.add_paragraph(statement, style="List Number")
            # add choices
            for opt in mcq.options:
                self.doc.add_paragraph(f"{opt}: {mcq.options[opt]}", style="List 2")
            # newline
            self.doc.add_paragraph()

        self.doc.save(self.docx_path)

    def add_ARS(self):
        """
        Add the valid Assertion Reasoning Questions to the .docx file
        """
        self.doc.add_heading("Assertion and Reasoning", level=1)

        AR.write_instructions(self.doc)

        for opt in AR.OPTIONS:
            self.doc.add_paragraph(f"{opt}: {AR.OPTIONS[opt]}", style="List 2")
        self.doc.add_paragraph()

        for ar in AR.QUESTIONS:
            self.attach_image(ar)

            assertion = ' '.join(ar.assertion)
            reason = ' '.join(ar.reason)

            paragraph = self.doc.add_paragraph(style="List Number")
            paragraph.add_run("Assertion (A): ").bold = True
            paragraph.add_run(assertion)

            paragraph = self.doc.add_paragraph(style="List 2")
            paragraph.add_run("Reason (R): ").bold = True
            paragraph.add_run(reason)
            paragraph.add_run('\n')

        self.doc.save(self.docx_path)

    def add_SUBS(self):
        """
        Add the valid Subjective Questions to the .docx file
        """
        self.doc.add_heading("Subjective Questions", level=1)

        SUB.write_instructions(self.doc)

        for sub in SUB.QUESTIONS:
            self.attach_image(sub)

            statement = ' '.join(sub.question)
            self.doc.add_paragraph(statement, style="List Number").add_run('\n')
        
        self.doc.save(self.docx_path)
