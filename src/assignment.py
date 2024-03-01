import docx

from src.question_classes import *
from src.utils import *

class Assignment:

    def __init__(self, path: str) -> None:
        """
        Initialize an Assignment object if `path` represents a valid .txt file
        """
        assert path.endswith(".txt"), "File must be of .txt format"
        try:
            file = open(path)
        except FileNotFoundError:
            print(f"File {path} was not found.")
            return
        
        # strip whitespaces from the line if it is not a newline
        contents = [line.strip('\n').strip() if line != '\n' else '\n' for line in file]
        file.close()

        # strip newlines from the start and end of contents
        start = 0
        while contents[start] == '\n':
            start += 1
        
        end = len(contents) - 1
        while contents[end] == '\n':
            end -=1
        # split the questions
        contents = split_contents(contents[start:end+1])
        # stores the questions as class objects
        self.QUESTIONS = {"MCQS": [], "ARS": [], "SUBS": [], "INVALIDATED": []}
        # process each question
        for question in contents:

            match question[0]:

                case 'MCQ':
                    self.QUESTIONS["MCQS"].append(MCQ.FromListContent(question[1:]))
                case 'AR':
                    self.QUESTIONS["ARS"].append(AR.FromListContent(question[1:]))
                case 'SUB':
                    self.QUESTIONS["SUBS"].append(SUB(question[1:]))
                case _:
                    self.QUESTIONS["INVALIDATED"].append(question)
        
        self.doc = docx.Document()


    def InitDoc(self):
        """
        Basic initializations for the .docx file
        """
        self.doc.add_heading("Assignment", level=0)
        self.doc.save("assignment.docx")
    

    def AddMCQS(self):
        """
        Add the valid Multiple Choice Questions to the .docx file
        """
        self.doc.add_heading("Multiple Choice Questions", level=1)
        self.doc.add_paragraph()
        for mcq in self.QUESTIONS["MCQS"]:
            # add statement
            statement = ' '.join(mcq.question)
            self.doc.add_paragraph(statement, style="List Number")
            # add choices
            for opt in mcq.options:
                self.doc.add_paragraph(f"{opt}: {mcq.options[opt]}", style="List 2")
            # newline
            self.doc.add_paragraph()

        self.doc.save("assignment.docx")


    def AddARS(self):
        """
        Add the valid Assertion Reasoning Questions to the .docx file
        """
        self.doc.add_heading("Assertion and Reasoning", level=1)
        # add instruction
        instruction = self.doc.add_paragraph()
        instruction.add_run("The following questions consist of two statements- Assertion (A) and Reasoning (R)\n")
        instruction.add_run("Answer these questions selecting the most appropriate option given below:")
        # add choices
        for opt in AR.OPTIONS:
            self.doc.add_paragraph(f"{opt}: {AR.OPTIONS[opt]}", style="List 2")
        self.doc.add_paragraph() # newline

        for ar in self.QUESTIONS["ARS"]:
            # add statements
            assertion = ' '.join(ar.assertion)
            reason = ' '.join(ar.reason)
            paragraph = self.doc.add_paragraph(style="List Number")
            paragraph.add_run("Assertion (A): ").bold = True
            paragraph.add_run(assertion)
            paragraph = self.doc.add_paragraph(style="List 2")
            paragraph.add_run("Reason (R): ").bold = True
            paragraph.add_run(reason)
            self.doc.add_paragraph()

        self.doc.save("assignment.docx")


    def AddSUBS(self):
        """
        Add the valid Subjective Questions to the .docx file
        """
        self.doc.add_heading("Subjective Questions", level=1)
        # add statement
        for sub in self.QUESTIONS["SUBS"]:
            statement = ' '.join(sub.question)
            self.doc.add_paragraph(statement, style="List Number")
            self.doc.add_paragraph()
        
        self.doc.save("assignment.docx")